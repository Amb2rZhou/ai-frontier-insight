"""Generate sources-and-prompts.docx"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.3

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()
    return t

# ═══════════════════════════════════════════════════
# Title
# ═══════════════════════════════════════════════════
title = doc.add_heading('AI Frontier Insight Bot', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph('数据源、筛选逻辑与 Prompt 设计')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
doc.add_paragraph()

# ═══════════════════════════════════════════════════
# Part 1: 数据源
# ═══════════════════════════════════════════════════
doc.add_heading('一、数据源', level=1)

# 1.1 X/Twitter
doc.add_heading('1. X/Twitter（~200 账号）', level=2)
doc.add_paragraph(
    '通过 X List 监控 AI 领域关键人物的推文。来源池约 300 人（via mitbunny.ai 数据库），'
    '经筛选（粉丝 ≥5,000，角色为 Research/Scientist/Engineer/Professor/Founder/CEO/CTO/Director/Lead）'
    '后同步约 200 人到 X List。过去一周活跃账号约 116 个，日均采集 ~90 条推文。'
)
doc.add_paragraph('账号按机构/身份分布：')
add_table(
    ['类别', '数量', '代表人物'],
    [
        ['OpenAI', '~30', 'Sam Altman, Mark Chen 等'],
        ['Google / DeepMind', '~32', 'Jeff Dean, Demis Hassabis 等'],
        ['NVIDIA', '~39', 'Jim Fan, Jensen Huang 等'],
        ['Anthropic', '~9', 'Dario Amodei, Amanda Askell 等'],
        ['Microsoft', '~6', '研究/产品团队'],
        ['Meta', '~2', 'Yann LeCun 等'],
        ['HuggingFace', '~15', '模型/平台团队成员'],
        ['Academia (Stanford/CMU/MIT/Harvard/Berkeley)', '~36', '各校 AI 教授'],
        ['Founders/CEOs', '~21', '各 AI 创业公司创始人'],
        ['Independent Researchers', '~19', 'Andrej Karpathy 等'],
        ['Other (博主/工程师/生态)', '~92', 'Swyx, Simon Willison 等'],
    ]
)
doc.add_paragraph('账号按权威性分 3 级（Tier 1 > 2 > 3），影响 signal 评分权重。')

# 1.2 RSS
doc.add_heading('2. RSS Feeds（28 个源）', level=2)
add_table(
    ['分组', '源'],
    [
        ['付费科技媒体', 'The Information'],
        ['英文科技媒体', 'TechCrunch, The Verge, Ars Technica, Wired, VentureBeat, MIT Technology Review'],
        ['AI 公司博客', 'OpenAI Blog, Google AI Blog, Meta Engineering, Anthropic Blog'],
        ['技术社区', 'Hacker News, Reddit ML, Reddit AI'],
        ['ML 平台', 'HuggingFace Blog'],
        ['AI Safety', 'LessWrong (Curated), Alignment Forum'],
        ['VC 洞察', 'a16z Podcast, a16z 16 Minutes, a16z Live'],
        ['中文播客', '硅谷101, 海外独角兽'],
        ['研究博客', "Lil'Log (Lilian Weng)"],
        ['Benchmark', 'LMSYS Arena Blog, Scale AI Blog, Aider Leaderboard, ARC Prize'],
    ]
)

# 1.3 GitHub
doc.add_heading('3. GitHub', level=2)
p = doc.add_paragraph()
p.add_run('Trending: ').bold = True
p.add_run('每日搜索 6 个 topic（artificial-intelligence, machine-learning, llm, agents, deep-learning, generative-ai）的热门仓库，取 top 30。')
p = doc.add_paragraph()
p.add_run('Watch Repos: ').bold = True
p.add_run('监控 13 个核心仓库的 release — openai-python, anthropic-sdk-python, LangChain, LangGraph, AutoGen, CrewAI, Transformers, vLLM, Ollama, OpenHands, browser-use, AutoGPT, SWE-bench。')

# 1.4 Arxiv
doc.add_heading('4. Arxiv 论文', level=2)
doc.add_paragraph(
    '通过 HuggingFace Daily Papers API 获取经社区投票筛选的论文（≥5 upvotes），'
    '覆盖 cs.AI / cs.CL / cs.LG / cs.CV / cs.MA 五个类别，取 top 25。'
    '相比原始 Arxiv API，信噪比大幅提升。'
)

# 1.5 HuggingFace
doc.add_heading('5. HuggingFace 模型 & Spaces', level=2)
doc.add_paragraph('Trending 模型 top 15 · 新发布模型 top 10 · Trending Spaces top 10')

# 1.6 Benchmark
doc.add_heading('6. Benchmark 排行榜（5 个）', level=2)
doc.add_paragraph('自动监控排行榜变动（榜首更替、新模型上榜等），重大变化纳入日报。')
add_table(
    ['排行榜', '衡量能力'],
    [
        ['Open LLM Leaderboard', '开源模型综合能力'],
        ['SWE-bench Verified', '代码工程能力'],
        ['ARC-AGI-2', '抽象推理能力'],
        ['OSWorld', '电脑操控能力'],
        ['Terminal-Bench 2.0', '终端编程能力'],
    ]
)

# ═══════════════════════════════════════════════════
# Part 2: 筛选流程
# ═══════════════════════════════════════════════════
doc.add_heading('二、筛选流程', level=1)

steps = [
    ('采集', '350+ 条/天原始信息（来自以上 6 类数据源）'),
    ('Signal Extraction', 'AI 模型从原始信息中提取 15 条候选信号，按权威性、互动量、新颖度评分'),
    ('代码去重', '与近 3 天已报道的信号标题对比，60% 词重叠即过滤'),
    ('排序截取', '按 signal_strength 降序排列，取 top 10'),
    ('Insight Generation', 'AI 模型为每条信号生成 insight（为什么重要）+ implication（意味着什么）'),
    ('Trend Update', '更新趋势数据库（accelerating / stable / decelerating / fading）'),
    ('Trend Summary', '识别跨源模式，输出 3-5 条核心趋势观察'),
    ('输出', '最终日报：10 条信号 + 洞察 + 趋势总结'),
]
add_table(
    ['步骤', '操作', '说明'],
    [[str(i+1), s[0], s[1]] for i, s in enumerate(steps)]
)

# ═══════════════════════════════════════════════════
# Part 3: Prompt 原文
# ═══════════════════════════════════════════════════
doc.add_heading('三、Prompt 原文', level=1)

def add_prompt_block(title, model_note, content):
    """Add a prompt section with title, model note, and raw content in code-style."""
    doc.add_heading(title, level=2)
    p = doc.add_paragraph()
    p.add_run('模型: ').bold = True
    p.add_run(model_note)
    doc.add_paragraph()
    # Add prompt content as monospace paragraphs
    for line in content.split('\n'):
        p = doc.add_paragraph(line)
        for run in p.runs:
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(0)
    doc.add_paragraph()  # spacing after block

# Read prompt files
prompt_files = {
    'signal_extraction': ('Prompt 1: Signal Extraction（信号提取）', 'DeepSeek（Haiku 级别，高吞吐低成本）'),
    'insight_generation': ('Prompt 2: Insight Generation（洞察生成）', 'DeepSeek（Sonnet 级别，深度分析）'),
    'trend_update': ('Prompt 3: Trend Update（趋势更新）', 'DeepSeek'),
    'trend_summary': ('Prompt 4: Trend Summary（趋势总结）', 'DeepSeek'),
    'weekly_deep_insight': ('Prompt 5: Weekly Deep Insight（周报深度分析）', 'DeepSeek（Sonnet 级别）'),
}

import os
for fname, (title, model) in prompt_files.items():
    path = f'/Users/zhouzhile/ai-frontier-insight/prompts/{fname}.txt'
    with open(path, 'r') as f:
        content = f.read().strip()
    add_prompt_block(title, model, content)

# ── Save ──
out = '/Users/zhouzhile/ai-frontier-insight/docs/AI Frontier Insight Bot - 数据源与筛选逻辑.docx'
doc.save(out)
print(f'Saved to: {out}')
